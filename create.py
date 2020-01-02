# -*- coding: cp1252 -*-
from docx import Document
import read_docx
import re
from datetime import datetime
import sys
droppedFile = sys.argv[1]
print droppedFile

months = (
        'January', 'February', 'March', 'April', 'May', 'June', 'July', 'August',
        'September', 'October', 'November', 'December'
         )

document = Document()

# Document import and parsing

records = []
total_letters = 0
docx_text = read_docx.get_docx_text('IMPORT.docx').splitlines()

total_letters = 0
date = ""
link = ""
letter_count = 0

def addEntry():
    if link != "":
        global date
        global link
        global letter_count
        global total_letters
        total_letters += letter_count
        records.append((date, link, letter_count))
        date = ""
        link = ""
        letter_count = 0


for line in docx_text:
    if line == "":
        continue
    
    date_match = re.search(r'^\d de \w', line)
    link_match = re.search(r'^http', line)

    if date_match:
        addEntry()
        date = line
    elif link_match:
        addEntry()
        link = line
    else:
        letter_count += len(line)
                
#From the results, generate the table and the letter count

p = document.add_paragraph('Total letras: ' + str(total_letters))
table = document.add_table(rows=1, cols=3)
table.allow_autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Data'
hdr_cells[1].text = 'Link'
hdr_cells[2].text = 'Carateres'
for date, link, letter_count in records:
    row_cells = table.add_row().cells
    row_cells[0].text = date
    row_cells[1].text = link
    row_cells[2].text = str(letter_count)

time = str(datetime.now().strftime('%Y-%m-%d-%H%M'))
document.save('Report-' + time + '.docx')

    

